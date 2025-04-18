
import streamlit.components.v1 as components


def inject_autoload_receiver(user_id):
    import streamlit.components.v1 as components
    js_code = f"""
    <script>
    window.addEventListener("message", (event) => {{
        if (event.data && event.data.type === "gost-autoload") {{
            const payload = event.data.payload;
            const input = window.parent.document.createElement("input");
            input.type = "hidden";
            input.name = "gost_autoload_data";
            input.value = JSON.stringify(payload);
            window.parent.document.body.appendChild(input);

            // Эмулируем изменение, чтобы Streamlit подхватил
            const evt = new Event("input", {{ bubbles: true }});
            input.dispatchEvent(evt);
        }}
    }});
    </script>
    """
    components.html(js_code, height=0)


def autosave_to_localstorage(user_id, data):
    import json
    js_code = f"""
    <script>
    localStorage.setItem("autosave_" + {json.dumps(user_id)}, JSON.stringify({json.dumps(data)}));
    </script>
    """
    components.html(js_code, height=0)


def autoload_from_localstorage(user_id):
    js_code = f"""
    <script>
    const data = localStorage.getItem("autosave_" + {json.dumps(user_id)});
    if (data) {{
        const parsed = JSON.parse(data);
        window.postMessage({{ type: "gost-autoload", payload: parsed }}, "*");
    }}
    </script>
    """
    components.html(js_code, height=0)




from streamlit_js_eval import streamlit_js_eval

def load_from_localstorage_with_js_eval(user_id):
    key = f"autosave_{user_id}"
    result = streamlit_js_eval(js_expressions=f"localStorage.getItem('{key}')", key="read_localstorage")
    return result

import streamlit as st
import os
import json
import requests


# Проверка наличия секретов

GITHUB_TOKEN = "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"
GITHUB_USERNAME = "islamshovgenov"
GITHUB_REPO = "gost-citations-app"
GITHUB_BRANCH = "main"



import sys
import streamlit as st
import re
import json
import os
import logging
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from io import BytesIO

st.set_page_config(page_title="Объединение ссылок по ГОСТ", layout="wide")

# spaCy временно отключён
nlp = None

# Настройка логирования
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Константы и пути
PROJECT_DIR = "projects"
AUTOSAVE_FILE = "project_autosave.json"
os.makedirs(PROJECT_DIR, exist_ok=True)

# Предкомпилированные регулярные выражения (используются в process_fragment и в случае отсутствия spaCy)
REFS_SPLIT_REGEX = re.compile(
    r"(?=^\s*(?:\[\d+\]|(?:19|20)\d{2}\.|[A-ZА-ЯЁ][a-zа-яё]+,|\d+\.))", 
    flags=re.M
)
CITE_REGEX = re.compile(r"\[(\d+)\]")

#########################################
# Функции для управления файлами проекта
#########################################
def load_project(project_path: str) -> dict:
    """Загружает проект из JSON файла."""
    try:
        with open(project_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return data
    except Exception as e:
        logging.error(f"Ошибка загрузки проекта: {e}")
        return {}

def save_project(project_path: str, data: dict) -> None:
    """Сохраняет проект в JSON файл."""
    try:
        with open(project_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logging.error(f"Ошибка сохранения проекта: {e}")

def autosave_project(data: dict) -> None:
    """Автосохранение проекта в локальный файл."""
    try:
        with open(AUTOSAVE_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logging.error(f"Ошибка автосохранения проекта: {e}")

#########################################
# Функция для разбора и очистки списка литературы с помощью spaCy (вариант 2)
#########################################
def parse_references(refs_text: str) -> dict:
    """
    Универсальный разбор списка литературы:
    - поддерживает нумерованные и не нумерованные списки;
    - работает с русскими и английскими записями;
    - определяет начало записи по различным шаблонам.
    """
    import re
    lines = refs_text.strip().split("\n")

    # Список возможных шаблонов начала новой записи
    patterns = [
        r"(?=^\s*\[\d+\])",                      # [1]
        r"(?=^\s*\d+[\.)])",                     # 1. или 1)
        r"(?=^\s*(19|20)\d{2}\.)",               # 2020. Текст...
        r"(?=^[A-ZА-ЯЁ][a-zа-яё]+.*?\(\d{4}\))", # Smith J. (2020) или Иванов И.И. (2020)
        r"(?=^[A-ZА-ЯЁ][a-zа-яё]+,)"             # Ivanov, A.
    ]

    best_split = []
    for pat in patterns:
        split_regex = re.compile(pat, flags=re.M)
        parts = [s.strip() for s in split_regex.split(refs_text) if s.strip()]
        if len(parts) > len(best_split):
            best_split = parts

    refs_dict = {}
    for idx, ref in enumerate(best_split, start=1):
        ref = ref.strip()
        if not ref:
            continue
        refs_dict[idx] = ref
    return refs_dict
from rapidfuzz import fuzz

def process_fragment(frag: dict, global_ref_map: dict, current_index: list) -> tuple:
    """
    Обрабатывает фрагмент, заменяя локальные ссылки на глобальные номера.
    Убирает дубликаты номеров и очищает [x] в начале ссылок литературы.
    """
    from rapidfuzz import fuzz
    local_refs_dict = frag['refs']
    refs_list = list(local_refs_dict.values())
    new_id = 1
    threshold = 90

    def clean_reference(text):
        return re.sub(r'^\[\d+\]\s*', '', text).strip()

    def find_existing_ref(new_ref_text):
        for known_text in global_ref_map:
            if fuzz.ratio(known_text.lower(), new_ref_text.lower()) >= threshold:
                return known_text
        return None

    def replace_cite(match):
        raw_num = int(match.group(1))
        if raw_num < 1 or raw_num > len(refs_list):
            return '[??]'
        raw_ref_text = refs_list[raw_num - 1]
        ref_text = clean_reference(raw_ref_text)

        existing_ref = find_existing_ref(ref_text)
        if existing_ref:
            ref_text = existing_ref
        else:
            global_ref_map[ref_text] = current_index[0]
            current_index[0] += 1

        return f'[{global_ref_map[ref_text]}]'

    frag_text = re.sub(r'\[(\d+)\]', replace_cite, frag['text'])
    return frag_text, global_ref_map, current_index
    def find_existing_ref(new_ref_text):
        for known_text in global_ref_map:
            if fuzz.ratio(known_text.lower(), new_ref_text.lower()) >= threshold:
                return known_text
        return None

    def replace_cite(match: re.Match) -> str:
        nonlocal new_id
        raw_num = int(match.group(1))
        if raw_num < 1 or raw_num > len(refs_list):
            logging.warning('Оригинальный номер ссылки не найден в списке литературы')
            return '[??]'
        ref_text = refs_list[raw_num - 1]

        if ref_text in local_text_map:
            local_num = local_text_map[ref_text]
        else:
            local_text_map[ref_text] = new_id
            local_num = new_id
            new_id += 1

        existing_ref = find_existing_ref(ref_text)
        if existing_ref:
            ref_text = existing_ref
        else:
            global_ref_map[ref_text] = current_index[0]
            current_index[0] += 1

        return f'[{global_ref_map[ref_text]}]'

    frag_text = re.sub(r'\[(\d+)\]', replace_cite, frag['text'])
    return frag_text, global_ref_map, current_index
    def replace_cite(match: re.Match) -> str:
        nonlocal new_id
        raw_num = int(match.group(1))
        # Проверяем, что номер есть в диапазоне ссылок (нумерация в списке начинается с 1)
        if raw_num < 1 or raw_num > len(local_refs):
            logging.warning("Оригинальный номер ссылки не найден в списке литературы")
            return "[??]"
        ref_text = local_refs[raw_num - 1]
        # Если ссылка уже встречалась в этом фрагменте, используем её локальный номер
        if ref_text in local_text_map:
            local_num = local_text_map[ref_text]
        else:
            local_text_map[ref_text] = new_id
            local_num = new_id
            new_id += 1
        # В глобальной карте, если ссылка новая, присваиваем ей глобальный номер
        if ref_text not in global_ref_map:
            global_ref_map[ref_text] = current_index[0]
            current_index[0] += 1
        return f"[{global_ref_map[ref_text]}]"

    frag_text = re.sub(r"\[(\d+)\]", replace_cite, frag["text"])
    return frag_text, global_ref_map, current_index
#########################################
# Инициализация переменных сессии
#########################################
def init_session_state(user_id):
    if f"{user_id}_fragments" not in st.session_state:
        st.session_state[f"{user_id}_fragments"] = []
    if f"{user_id}_ref_map" not in st.session_state:
        st.session_state[f"{user_id}_ref_map"] = {}
    if f"{user_id}_ref_counter" not in st.session_state:
        st.session_state[f"{user_id}_ref_counter"] = 1
    if f"{user_id}_final_text" not in st.session_state:
        st.session_state[f"{user_id}_final_text"] = ""
    if f"{user_id}_final_refs" not in st.session_state:
        st.session_state[f"{user_id}_final_refs"] = []
    if f"{user_id}_fragments" not in st.session_state:
        st.session_state[f"{user_id}_fragments"] = []
    if 'ref_map' not in st.session_state:
        st.session_state[f"{user_id}_ref_map"] = {}
    if 'ref_counter' not in st.session_state:
        st.session_state[f"{user_id}_ref_counter"] = 1
    if 'final_text' not in st.session_state:
        st.session_state[f"{user_id}_final_text"] = ""
    if 'final_refs' not in st.session_state:
        st.session_state[f"{user_id}_final_refs"] = []
    if 'edit_index' not in st.session_state:
        st.session_state.edit_index = None
    if 'start_index' not in st.session_state:
        st.session_state.start_index = 1

def restore_autosave(user_id):
    if 'restored' not in st.session_state:
        if os.path.exists(AUTOSAVE_FILE):
            try:
                with open(AUTOSAVE_FILE, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    st.session_state[f"{user_id}_fragments"] = data.get("fragments", [])
                    st.session_state[f"{user_id}_ref_map"] = data.get("ref_map", {})
                    st.session_state[f"{user_id}_ref_counter"] = data.get("ref_counter", 1)
                    st.session_state[f"{user_id}_final_text"] = data.get("final_text", "")
                    st.session_state[f"{user_id}_final_refs"] = data.get("final_refs", [])
                st.success("✅ Автоматически восстановлен последний проект")
            except Exception as e:
                st.warning(f"Ошибка при восстановлении: {e}")
        st.session_state.restored = True

def update_autosave():
    data = {
        "fragments": st.session_state.get("fragments", []),
        "ref_map": st.session_state.get("ref_map", {}),
        "ref_counter": st.session_state.get("ref_counter", 1),
        "final_text": st.session_state.get("final_text", ""),
        "final_refs": st.session_state.get("final_refs", [])
    }
    autosave_project(data)

#########################################
# Основная функция приложения
#########################################
from docx import Document
from docx.shared import Pt

def generate_docx(text, references):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    doc.add_paragraph(text.strip())
    doc.add_paragraph("\nСписок литературы:")
    for ref in references:
        doc.add_paragraph(ref, style='List Number')

    return doc

def main():

    # Автоматическая загрузка последнего проекта
    if "last_opened_project" in st.session_state:
        last_proj = os.path.join(PROJECT_DIR, st.session_state["last_opened_project"])
        if os.path.exists(last_proj):
            data = load_project(last_proj)
            st.session_state.fragments = data.get("fragments", [])
            st.session_state.ref_map = data.get("ref_map", {})
            st.session_state.ref_counter = data.get("ref_counter", 1)
            st.session_state.final_text = data.get("final_text", "")
            st.session_state.final_refs = data.get("final_refs", [])
            st.session_state.restored = True

    user_id = st.sidebar.text_input("🧙 Ваше имя, мудрейший из оформителей ГОСТа", value="Безымянный")
    init_session_state(user_id)
    restore_autosave(user_id)
    
    #########################################
    # Просмотр добавленных фрагментов с возможностью редактирования и удаления
    #########################################
    st.subheader("📋 Добавленные фрагменты")
    for idx, frag in enumerate(st.session_state[f"{user_id}_fragments"]):
        with st.expander(f"Фрагмент {idx + 1}", expanded=False):
            with st.form(key=f"fragment_form_{idx}"):
                st.markdown(f"**Текст:**\n{frag['text']}")
                st.markdown("**Список литературы:**")
                if isinstance(frag.get("refs"), dict):
                    for orig_num, ref_text in frag["refs"].items():
                        st.markdown(f"{orig_num}. {ref_text}")
                elif isinstance(frag.get("refs"), list):
                    for i, ref_text in enumerate(frag["refs"]):
                        st.markdown(f"{i+1}. {ref_text}")

                col1, col2, col3, col4 = st.columns([1, 1, 1, 2])
                with col1:
                    up = st.form_submit_button("⬆️")
                    if up and idx > 0:
                        st.session_state[f"{user_id}_fragments"][idx - 1], st.session_state[f"{user_id}_fragments"][idx] = \
                            st.session_state[f"{user_id}_fragments"][idx], st.session_state[f"{user_id}_fragments"][idx - 1]
                        st.rerun()
                with col2:
                    down = st.form_submit_button("⬇️")
                    if down and idx < len(st.session_state[f"{user_id}_fragments"]) - 1:
                        st.session_state[f"{user_id}_fragments"][idx + 1], st.session_state[f"{user_id}_fragments"][idx] = \
                            st.session_state[f"{user_id}_fragments"][idx], st.session_state[f"{user_id}_fragments"][idx + 1]
                        st.rerun()
                with col3:
                    edit_button = st.form_submit_button("✏️ Редактировать")
                    if edit_button:
                        st.session_state.edit_index = idx
                        st.rerun()
                with col4:
                    delete_button = st.form_submit_button("🗑 Удалить")
                    if delete_button:
                        st.session_state[f"{user_id}_fragments"].pop(idx)
                        st.rerun()
    # Статистика по проекту
    st.sidebar.markdown(f"**Фрагментов:** {len(st.session_state[f"{user_id}_fragments"])}")
    st.sidebar.markdown(f"**Итоговых ссылок:** {len(st.session_state[f"{user_id}_final_refs"])}")

    local_data = load_from_localstorage_with_js_eval(user_id)
    if local_data:
        try:
            payload = json.loads(local_data)
            st.session_state[f"{user_id}_fragments"] = payload.get("fragments", [])
            st.session_state[f"{user_id}_ref_map"] = payload.get("ref_map", {})
            st.session_state[f"{user_id}_ref_counter"] = payload.get("ref_counter", 1)
            st.session_state[f"{user_id}_final_text"] = payload.get("final_text", "")
            st.session_state[f"{user_id}_final_refs"] = payload.get("final_refs", [])
            st.session_state.restored = True
            st.success("✅ Проект восстановлен из LocalStorage (через JS Eval)")
        except Exception as e:
            st.warning(f"Ошибка парсинга автосохранения: {e}")

    autoload_from_localstorage(user_id)
    inject_autoload_receiver(user_id)
    st.text_input("hidden_autoload_input", value="", key="gost_autoload_data", label_visibility="collapsed")

    # Обработка автозагруженных данных
if "gost_autoload_data" in st.session_state:
    raw_data = st.session_state["gost_autoload_data"]
    if raw_data:
        try:
            payload = json.loads(raw_data)
            st.session_state[f"{user_id}_fragments"] = payload.get("fragments", [])
            st.session_state[f"{user_id}_ref_map"] = payload.get("ref_map", {})
            st.session_state[f"{user_id}_ref_counter"] = payload.get("ref_counter", 1)
            st.session_state[f"{user_id}_final_text"] = payload.get("final_text", "")
            st.session_state[f"{user_id}_final_refs"] = payload.get("final_refs", [])
            st.session_state.restored = True
            st.success("✅ Проект восстановлен из LocalStorage!")
        except Exception as e:
            st.warning(f"Ошибка загрузки из LocalStorage: {e}")
    else:
        st.info("ℹ️ LocalStorage пуст — нет данных для восстановления.")
    st.title("Автоматическое объединение ссылок и списка литературы (ГОСТ)")



    # Панель управления проектами
    st.sidebar.title("📁 Управление проектами")
    projects_files = [p for p in os.listdir(PROJECT_DIR) if p.endswith(".json")]
    chosen_file = st.sidebar.selectbox("Выбрать проект из списка", ["—"] + projects_files)
    if chosen_file != "—":
        st.session_state["last_opened_project"] = chosen_file  # Запоминаем последний проект
        project_name = chosen_file.replace(".json", "")
    else:
        project_name = st.sidebar.text_input("Или ввести название проекта вручную", value="default")
        project_path = os.path.join(PROJECT_DIR, f"{project_name}.json")
    
    # Кнопки сохранения, загрузки, удаления, импорта и экспорта проекта
    if st.sidebar.button("💾 Сохранить проект"):
        data_to_save = {
            "fragments": st.session_state.get("fragments", []),
            "ref_map": st.session_state.get("ref_map", {}),
            "ref_counter": st.session_state.get("ref_counter", 1),
            "final_text": st.session_state.get("final_text", ""),
            "final_refs": st.session_state.get("final_refs", [])
        }
        save_project(project_path, data_to_save)
        st.sidebar.success(f"Проект '{project_name}' сохранён")
    
    if st.sidebar.button("📂 Загрузить проект"):
        if os.path.exists(project_path):
            data = load_project(project_path)
            st.session_state[f"{user_id}_fragments"] = data.get("fragments", [])
            st.session_state[f"{user_id}_ref_map"] = data.get("ref_map", {})
            st.session_state[f"{user_id}_ref_counter"] = data.get("ref_counter", 1)
            st.session_state[f"{user_id}_final_text"] = data.get("final_text", "")
            st.session_state[f"{user_id}_final_refs"] = data.get("final_refs", [])
            st.sidebar.success(f"Проект '{project_name}' загружен")
        else:
            st.sidebar.error("Файл проекта не найден")
    
    if st.sidebar.button("🗑 Удалить проект"):
        if os.path.exists(project_path):
            os.remove(project_path)
            st.sidebar.success(f"Проект '{project_name}' удалён")
        else:
            st.sidebar.error("Такого проекта нет в папке")
    
    uploaded_file = st.sidebar.file_uploader("📥 Импорт из файла (.json)", type="json")
    if uploaded_file is not None:
        try:
            data = json.load(uploaded_file)
            st.session_state[f"{user_id}_fragments"] = data.get("fragments", [])
            st.session_state[f"{user_id}_ref_map"] = data.get("ref_map", {})
            st.session_state[f"{user_id}_ref_counter"] = data.get("ref_counter", 1)
            st.session_state[f"{user_id}_final_text"] = data.get("final_text", "")
            st.session_state[f"{user_id}_final_refs"] = data.get("final_refs", [])
            st.sidebar.success("Проект импортирован из файла")
        except Exception as e:
            st.sidebar.error(f"Ошибка при импорте: {e}")
    
    if st.sidebar.button("📤 Экспорт в файл"):
        export_data = json.dumps({
            "fragments": st.session_state.get("fragments", []),
            "ref_map": st.session_state.get("ref_map", {}),
            "ref_counter": st.session_state.get("ref_counter", 1),
            "final_text": st.session_state.get("final_text", ""),
            "final_refs": st.session_state.get("final_refs", [])
        }, ensure_ascii=False, indent=2)
        st.sidebar.download_button(
            "📎 Скачать JSON",
            export_data,
            file_name=f"{project_name}.json",
            mime="application/json"
        )
    
    st.sidebar.markdown("---")
    st.sidebar.subheader("⚙️ Настройки")
    st.session_state.start_index = st.sidebar.number_input("Начальный номер глобальной нумерации", min_value=1, value=1)
    st.sidebar.markdown("---")
    
    # Автовосстановление последнего проекта
    if "last_opened_project" in st.session_state and st.session_state["last_opened_project"].endswith(".json"):
        default_project = os.path.join(PROJECT_DIR, st.session_state["last_opened_project"])
        if os.path.exists(default_project):
            data = load_project(default_project)
            st.session_state.fragments = data.get("fragments", [])
            st.session_state.ref_map = data.get("ref_map", {})
            st.session_state.ref_counter = data.get("ref_counter", 1)
            st.session_state.final_text = data.get("final_text", "")
            st.session_state.final_refs = data.get("final_refs", [])
            st.success(f"Проект {st.session_state['last_opened_project']} автоматически восстановлен")
    
    # Инструкция для пользователя
    st.markdown("""
    #### 📌 Инструкция:
    - Вставляйте **текст фрагмента**, где уже есть ссылки вида [1], [2] и т.д.
    - Вставьте список литературы к этому фрагменту (в порядке ссылок).
    - Поддерживаются фрагменты, где ссылки начинаются не с [1] — нумерация будет скорректирована.
    - Повторяющиеся ссылки (по точному тексту) объединяются с сохранением одного номера.
    """)
    
    #########################################
    # Форма добавления / редактирования фрагмента
    #########################################
    with st.form(key="fragment_form", clear_on_submit=True):
        default_text = ""
        default_refs = ""
        if st.session_state.edit_index is not None:
            frag = st.session_state[f"{user_id}_fragments"][st.session_state.edit_index]
            default_text = frag["text"]
            if isinstance(frag.get("refs"), dict):
                default_refs = "\n\n".join(f"{k}. {v}" for k, v in frag["refs"].items())
            elif isinstance(frag.get("refs"), list):
                default_refs = "\n\n".join(f"{i+1}. {r}" for i, r in enumerate(frag["refs"]))
        new_text_input = st.text_area("Текст с локальной нумерацией (пример: [1], [2]...)", 
                                      value=default_text, height=200)
        new_refs_input = st.text_area("Список литературы к этому фрагменту", 
                                      value=default_refs, height=200)
        submitted = st.form_submit_button("💾 Сохранить фрагмент")
    
    if submitted and new_text_input.strip() and new_refs_input.strip():
        cleaned_refs = parse_references(new_refs_input.strip())
        fragment = {
            "text": new_text_input.strip(),
            "refs": cleaned_refs
        }
        if st.session_state.edit_index is not None:
            st.session_state[f"{user_id}_fragments"][st.session_state.edit_index] = fragment
            st.session_state.edit_index = None
        else:
            st.session_state[f"{user_id}_fragments"].append(fragment)
        update_autosave()
    #########################################
    # Объединение фрагментов и формирование итогового текста и списка литературы
    #########################################
    if st.button("📄 Объединить всё", key="combine_center"):
        new_text = ""
        new_refs = []
        global_ref_map = {}
        current_index = [st.session_state.get("start_index", 1)]
        all_issues = []

        for frag_idx, frag in enumerate(st.session_state[f"{user_id}_fragments"]):
            cited_numbers = set(int(n) for n in re.findall(r"\[(\d+)\]", frag["text"]))
            # Приведение ключей к int (если str)
            available_numbers = set(int(k) for k in frag["refs"].keys())

            missing = sorted(cited_numbers - available_numbers)
            extra = sorted(available_numbers - cited_numbers)

            if missing:
                all_issues.append(f"⚠️ Фрагмент {frag_idx+1}: ссылки в тексте, которых нет в списке — {missing}")

            frag_text, global_ref_map, current_index = process_fragment(frag, global_ref_map, current_index)
            frag_text = re.sub(r"\[\?\?\]", ":red[??]", frag_text)
            new_text += frag_text + "\n\n"

        if all_issues:
            st.warning("🚨 Обнаружены проблемы при объединении:")
            for issue in all_issues:
                st.markdown(issue)

        sorted_refs = sorted(global_ref_map.items(), key=lambda x: x[1])
        for ref_text, ref_num in sorted_refs:
            new_refs.append(f"[{ref_num}] {ref_text}")

        st.session_state[f"{user_id}_final_text"] = new_text
        st.session_state[f"{user_id}_final_refs"] = new_refs
        st.session_state[f"{user_id}_ref_map"] = global_ref_map
        st.session_state[f"{user_id}_ref_counter"] = current_index[0]
        update_autosave()

        st.subheader("📄 Объединённый текст")
        st.code(new_text, language="markdown")
        st.subheader("📚 Общий список литературы")
        for ref in new_refs:
            st.markdown(ref)
        st.success("Фрагменты объединены с учётом повторяющихся ссылок")
        for ref in new_refs:
            st.markdown(ref)
        st.success("Фрагменты объединены с учётом повторяющихся ссылок")
    
    #########################################
    # Вывод итогового результата и экспорт в DOCX
    #########################################
    if st.session_state[f"{user_id}_final_text"]:
        st.markdown("---")
        st.code(st.session_state[f"{user_id}_final_text"].strip(), language="markdown")
    
        for ref in st.session_state[f"{user_id}_final_refs"]:
            st.markdown(ref)
    
        if st.button("📥 Скачать DOCX", key="download_docx"):
            doc = Document()
            style = doc.styles["Normal"]
            font = style.font
            font.name = "Times New Roman"
            font.size = Pt(14)
            rFonts = style.element.rPr.rFonts
            rFonts.set(qn("w:eastAsia"), "Times New Roman")
    
            doc.add_paragraph("Текст обзора:")
            for paragraph in st.session_state[f"{user_id}_final_text"].strip().split("\n"):
                doc.add_paragraph(paragraph)
    
            doc.add_paragraph("\nСписок литературы:")
            for ref in st.session_state[f"{user_id}_final_refs"]:
                doc.add_paragraph(ref)
    
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
    
            st.download_button(
                label="📥 Скачать DOCX файл",
                data=buffer,
                file_name="обзор_со_ссылками.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_docx_file"
            )


    if st.session_state.get(f"{user_id}_final_text", ""):
        st.subheader("📤 Экспорт в DOCX")
        docx = generate_docx(
            st.session_state.get(f"{user_id}_final_text", ""),
            st.session_state.get(f"{user_id}_final_refs", [])
        )
        buffer = BytesIO()
        docx.save(buffer)
        buffer.seek(0)
        st.download_button(
            "📥 Скачать DOCX",
            buffer,
            file_name="citations.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
if __name__ == "__main__":
    main()
